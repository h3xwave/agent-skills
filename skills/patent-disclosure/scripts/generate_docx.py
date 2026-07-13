#!/usr/bin/env python3
import re
import sys
from pathlib import Path

_DOCX_IMPORT_ERROR = None
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt, RGBColor
except ImportError as exc:  # Keep CLI errors explicit when python-docx is unavailable.
    _DOCX_IMPORT_ERROR = exc
    Document = None


MARKDOWN_HEADING = re.compile(r"^(#{1,6})\s+(.+)$")
NUMERIC_HEADING = re.compile(r"^\d+(?:\.\d+)*\s+.+$")
TABLE_SEPARATOR = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)


def set_font(run, font_name="楷体_GB2312", size=12, bold=False, color="000000"):
    """Set East Asian and Western font styles for a text run."""
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        try:
            run.font.size = Pt(float(size))
        except (TypeError, ValueError):
            pass
    run.font.bold = bold
    if isinstance(color, str) and re.fullmatch(r"[0-9A-Fa-f]{6}", color):
        run.font.color.rgb = RGBColor.from_string(color)


def is_horizontal_rule(raw_line):
    """Return True for Markdown or template separator lines."""
    return bool(re.fullmatch(r"(?:-{3,}|\*{3,}|_{3,}|━+)", raw_line))


def parse_list_item(raw_line):
    """Parse simple Markdown list items and return display prefix plus text."""
    unordered = re.match(r"^[-*+]\s+(.+)$", raw_line)
    if unordered:
        return "• ", unordered.group(1)

    ordered = re.match(r"^(\d+[.、])\s*(.+)$", raw_line)
    if ordered:
        return f"{ordered.group(1)} ", ordered.group(2)

    return None


def add_formatted_runs(paragraph, text, size=12):
    """Add runs with simple **bold** support."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size)


def _meaningful_lines(markdown_text):
    return [
        line.strip()
        for line in markdown_text.splitlines()
        if line.strip() and not is_horizontal_rule(line.strip())
    ]


def _title_from_line(raw_line):
    heading = MARKDOWN_HEADING.match(raw_line)
    if heading and len(heading.group(1)) == 1:
        return heading.group(2).strip()
    if (
        not raw_line.startswith(("#", "```", "|", "- ", "* ", "+ "))
        and not NUMERIC_HEADING.match(raw_line)
    ):
        return raw_line
    return None


def validate_markdown(markdown_text, output_path=None):
    """Return user-correctable input errors without creating a document."""
    errors = []
    meaningful = _meaningful_lines(markdown_text)
    if not meaningful:
        return ["输入 Markdown 为空。"]

    title = _title_from_line(meaningful[0])
    if not title:
        errors.append("第一条有效内容必须是 '# 发明名称' 或纯文本发明名称。")

    remaining = meaningful[1:]
    has_section = any(
        NUMERIC_HEADING.match(line) or MARKDOWN_HEADING.match(line)
        for line in remaining
    )
    if not has_section:
        errors.append("正文至少需要一个数字章节标题或 Markdown 标题。")

    if any(TABLE_SEPARATOR.match(line) for line in meaningful):
        errors.append("输入包含 Markdown 表格；当前导出脚本不支持表格渲染。")

    in_mermaid = False
    for line in meaningful:
        if not in_mermaid and line.lower().startswith("```mermaid"):
            in_mermaid = True
        elif in_mermaid and line.startswith("```"):
            in_mermaid = False
    if in_mermaid:
        errors.append("Mermaid 代码块缺少结束围栏 ```。")

    if output_path is not None:
        output = Path(output_path)
        if output.suffix.lower() != ".docx":
            errors.append("输出路径必须使用 .docx 扩展名。")
        if not output.parent.exists():
            errors.append(f"输出目录不存在：{output.parent}")

    return errors


def add_page_number(doc):
    """Add 'Page X of Y' to the footer."""
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    set_font(run, size=10)
    run.text = "第 "

    current_begin = OxmlElement("w:fldChar")
    current_begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._element.append(current_begin)
    current_instruction = OxmlElement("w:instrText")
    current_instruction.set(qn("xml:space"), "preserve")
    current_instruction.text = "PAGE"
    paragraph.add_run()._element.append(current_instruction)
    current_end = OxmlElement("w:fldChar")
    current_end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._element.append(current_end)

    run = paragraph.add_run()
    set_font(run, size=10)
    run.text = " 页 共 "

    total_begin = OxmlElement("w:fldChar")
    total_begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._element.append(total_begin)
    total_instruction = OxmlElement("w:instrText")
    total_instruction.set(qn("xml:space"), "preserve")
    total_instruction.text = "NUMPAGES"
    paragraph.add_run()._element.append(total_instruction)
    total_end = OxmlElement("w:fldChar")
    total_end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._element.append(total_end)

    run = paragraph.add_run()
    set_font(run, size=10)
    run.text = " 页"


def _add_title(doc, title):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    set_font(run, size=14, bold=True)


def _add_heading(doc, text, level):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, size=12, bold=True)


def generate_docx(markdown_text, output_path):
    """Validate Markdown and generate an A4 Word document."""
    errors = validate_markdown(markdown_text, output_path)
    if errors:
        raise ValueError("\n".join(errors))
    if _DOCX_IMPORT_ERROR is not None:
        raise RuntimeError(f"缺少 python-docx 依赖：{_DOCX_IMPORT_ERROR}")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1.5

    title_pending = True
    in_mermaid = False

    for line in markdown_text.splitlines():
        raw_line = line.strip()
        if not raw_line or is_horizontal_rule(raw_line):
            continue

        if title_pending:
            title = _title_from_line(raw_line)
            _add_title(doc, title)
            title_pending = False
            continue

        if raw_line.lower().startswith("```mermaid"):
            in_mermaid = True
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("[请在此处手动插入对应的流程图/原理图图片]")
            set_font(run, size=12, bold=True, color="0000FF")
            continue

        if in_mermaid:
            if raw_line.startswith("```"):
                in_mermaid = False
            continue

        markdown_heading = MARKDOWN_HEADING.match(raw_line)
        if markdown_heading:
            level = 1 if len(markdown_heading.group(1)) <= 2 else 2
            _add_heading(doc, markdown_heading.group(2).strip(), level)
            continue

        if re.match(r"^\d+\s+", raw_line):
            _add_heading(doc, raw_line, 1)
            continue

        if re.match(r"^\d+(?:\.\d+)+\s+", raw_line):
            _add_heading(doc, raw_line, 2)
            continue

        if re.match(r"^(日期|技术交底书撰写人|电话|邮箱)(?:（[^）]*）)?[：:]", raw_line):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(raw_line)
            set_font(run, size=12)
            continue

        list_item = parse_list_item(raw_line)
        if list_item:
            prefix, text = list_item
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(24)
            paragraph.paragraph_format.first_line_indent = Pt(-12)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(prefix)
            set_font(run, size=12)
            add_formatted_runs(paragraph, text)
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(24)
        add_formatted_runs(paragraph, raw_line)

    add_page_number(doc)
    doc.save(output_path)


def main(argv=None):
    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) != 2:
        print(
            "Usage: python generate_docx.py <input.md> <output.docx>",
            file=sys.stderr,
        )
        return 2

    markdown_file = Path(args[0])
    output_file = Path(args[1])
    if not markdown_file.is_file():
        print(f"Input error: file not found: {markdown_file}", file=sys.stderr)
        return 2

    try:
        markdown_text = markdown_file.read_text(encoding="utf-8-sig")
    except (OSError, UnicodeError) as exc:
        print(f"Runtime error: unable to read {markdown_file}: {exc}", file=sys.stderr)
        return 1

    if output_file.exists() and output_file.is_dir():
        print(f"Runtime error: output path is a directory: {output_file}", file=sys.stderr)
        return 1

    errors = validate_markdown(markdown_text, output_file)
    if errors:
        for error in errors:
            print(f"Input error: {error}", file=sys.stderr)
        return 2

    if _DOCX_IMPORT_ERROR is not None:
        print(f"Runtime error: missing python-docx dependency: {_DOCX_IMPORT_ERROR}", file=sys.stderr)
        return 1

    try:
        generate_docx(markdown_text, output_file)
    except ValueError as exc:
        print(f"Input error: {exc}", file=sys.stderr)
        return 2
    except Exception as exc:
        print(f"Runtime error: unable to create {output_file}: {exc}", file=sys.stderr)
        return 1

    if not output_file.is_file():
        print(f"Runtime error: output was not created: {output_file}", file=sys.stderr)
        return 1

    print(f"Success: Word document saved to {output_file}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
